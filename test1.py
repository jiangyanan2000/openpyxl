import docx
from docx import Document
from docx.shared import RGBColor, Inches,Pt,Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_UNDERLINE,WD_TAB_ALIGNMENT,WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
import os


doc = Document("12.docx")

doc1 = Document("13.docx")

# for style1 in doc1.styles:
#     print(style1)
# print(len(doc.styles))
# for style2 in doc.styles:
#     print(style2)
t = doc1.styles
# par_styles = [s for s in t if s.type == WD_STYLE_TYPE.PARAGRAPH]
# for style in par_styles:
#     print(style.name)
d = doc.styles
d1 = d.add_style("yangshi",WD_STYLE_TYPE.CHARACTER)
print(d1.name)


for para in doc.paragraphs:
    print(para.text)
# s = doc1.paragraphs[0]
#
s1 = doc.paragraphs[0]
# s.paragraph_format.space_after = Inches(0.5)
# print(dir(s))
r1 = s1.add_run("胶州市胶东街道办事处文件",style="yangshi")
# r = s.add_run("胶州市胶东街道办事处文件",style="yangshi")


# r.font.name = '方正小标宋_GBK'
# r._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')
# r.font.size = Pt(72)
# # r.font.bold = True
# r.font.color.rgb = RGBColor(255,0,0)


# print(doc.paragraphs[0].text)
# for i in range(len(doc.paragraphs)):
#     print(str(i),  doc.paragraphs[i].text)

doc.save("12.docx")
doc1.save("13.docx")